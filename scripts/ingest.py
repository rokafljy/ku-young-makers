"""data/raw/의 문서를 파싱해 청크·임베딩을 data/processed/에 저장.

실행:
    python scripts/ingest.py

GEMINI_API_KEY 환경변수가 필요합니다.
"""
import json
import sys
from pathlib import Path

import faiss
import numpy as np

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from src.llm import embed_texts  # noqa: E402
from scripts.extract_qa import is_qa_like, extract_pairs  # noqa: E402

RAW_DIR = ROOT / "data" / "raw"
OUT_DIR = ROOT / "data" / "processed"
CHUNKS_PATH = OUT_DIR / "chunks.jsonl"
INDEX_PATH = OUT_DIR / "index.faiss"

CHUNK_SIZE = 700
CHUNK_OVERLAP = 100
EMBED_BATCH = 50


def read_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            yield i, text


def read_docx(path):
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    if paragraphs:
        yield 1, "\n".join(paragraphs)


def read_text(path):
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if text:
        yield 1, text


def read_file(path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        yield from read_pdf(path)
    elif suffix == ".docx":
        yield from read_docx(path)
    elif suffix in (".txt", ".md"):
        yield from read_text(path)
    elif suffix == ".hwp":
        print(f"[경고] HWP는 직접 지원되지 않습니다. PDF로 변환 후 재시도하세요: {path.name}")
    else:
        print(f"[건너뜀] 지원하지 않는 형식: {path.name}")


def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    buf = ""
    for p in paragraphs:
        if not buf:
            buf = p
        elif len(buf) + len(p) + 1 <= size:
            buf = buf + "\n" + p
        else:
            chunks.append(buf)
            buf = p
        if len(buf) >= size:
            # flush and start with overlap tail
            chunks.append(buf[:size])
            buf = buf[max(0, size - overlap):]
    if buf:
        chunks.append(buf)
    # split any oversize chunks by sliding window
    out = []
    for ch in chunks:
        if len(ch) <= size:
            out.append(ch)
        else:
            step = max(1, size - overlap)
            for i in range(0, len(ch), step):
                out.append(ch[i:i + size])
    return [c for c in out if c.strip()]


def build_chunks():
    chunks = []
    cid = 0
    files = [p for p in sorted(RAW_DIR.rglob("*")) if p.is_file() and not p.name.startswith(".")]
    if not files:
        return chunks
    for path in files:
        print(f"[처리] {path.name}")
        pages = list(read_file(path))
        if not pages:
            continue
        full_text = "\n".join(t for _, t in pages)
        if is_qa_like(full_text):
            pairs = extract_pairs(full_text)
            print(f"  - Q&A 구조 감지: {len(pairs)}쌍")
            for pair in pairs:
                text = f"질문: {pair['question']}\n답변: {pair['answer']}"
                chunks.append({
                    "id": cid, "source": path.name, "page": None,
                    "text": text, "type": "qa",
                })
                cid += 1
        else:
            for page_num, page_text in pages:
                for piece in chunk_text(page_text):
                    chunks.append({
                        "id": cid, "source": path.name, "page": page_num,
                        "text": piece, "type": "doc",
                    })
                    cid += 1
    return chunks


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    chunks = build_chunks()
    if not chunks:
        print("[오류] data/raw/에 처리할 문서가 없습니다. PDF/DOCX를 업로드한 뒤 다시 실행하세요.")
        return 1
    print(f"[임베딩] 총 청크 {len(chunks)}개")
    vectors = []
    for i in range(0, len(chunks), EMBED_BATCH):
        batch = chunks[i:i + EMBED_BATCH]
        vecs = embed_texts([c["text"] for c in batch], task_type="retrieval_document")
        vectors.extend(vecs)
        print(f"  - 진행 {min(i + EMBED_BATCH, len(chunks))}/{len(chunks)}")
    arr = np.array(vectors, dtype="float32")
    faiss.normalize_L2(arr)
    index = faiss.IndexFlatIP(arr.shape[1])
    index.add(arr)
    faiss.write_index(index, str(INDEX_PATH))
    with CHUNKS_PATH.open("w", encoding="utf-8") as f:
        for c in chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")
    print(f"[완료] 저장: {CHUNKS_PATH.relative_to(ROOT)}, {INDEX_PATH.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
